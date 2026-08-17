"""
FastAPI backend for the AI Quiz Generator.
"""
import time
import uuid
import logging
import hashlib
import io
import random
import string
import json
import base64
import os
import tempfile
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Header, Depends, Request
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from typing import Optional, Dict, List, Any

# Image & OCR Libraries
from PIL import Image
import pytesseract

# Word Document Export Library
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# NEW: PDF Export Libraries
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors

# NOTE FOR WINDOWS: Point pytesseract to the installed executable
#pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Existing modules (UNTOUCHED)
import database
import auth
from document_extractor import get_document_page_count, extract_text_from_document
from youtube_extractor import get_youtube_transcript
from quiz_generator import generate_quiz_from_large_text
from grading_engine import check_mcq, check_fill_blank, grade_long_answer
from cache_manager import generate_hash, get_cached_quiz, save_quiz_to_cache

# ==========================================
# 📋 SYSTEM LOGGING SETUP
# ==========================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("app_system.log"), 
        logging.StreamHandler()                
    ]
)
logger = logging.getLogger(__name__)

app = FastAPI(title="AI Quiz Generator API", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,  # ✅ UPDATE: Isay False kar diya gaya hai
    allow_methods=["*"],
    allow_headers=["*"],
)

database.init_db()

# 🪄 Create Bookmarks table dynamically if it doesn't exist (Zero manual DB changes needed)
conn = database.get_db_connection()
conn.execute('''CREATE TABLE IF NOT EXISTS bookmarked_questions (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_id INTEGER,
    quiz_id INTEGER,
    question_type TEXT,
    question_data TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)''')
conn.commit()
conn.close()

# ==========================================
# 🛡️ MIDDLEWARE: Request Tracking ID
# ==========================================
@app.middleware("http")
async def log_requests(request: Request, call_next):
    req_id = str(uuid.uuid4())[:8]
    request.state.req_id = req_id 
    logger.info(f"[ReqID: {req_id}] STARTED: {request.method} {request.url.path}")
    start_time = time.time()
    try:
        response = await call_next(request)
        process_time = time.time() - start_time
        logger.info(f"[ReqID: {req_id}] COMPLETED: Status {response.status_code} in {process_time:.2f}s")
        response.headers["X-Request-ID"] = req_id
        return response
    except Exception as e:
        logger.error(f"[ReqID: {req_id}] CRASHED: {str(e)}")
        raise e

# --- SCHEMAS ---
class RegisterRequest(BaseModel):
    name: str
    email: EmailStr
    password: str

class LoginRequest(BaseModel):
    email: EmailStr
    password: str

class UpdateProfileRequest(BaseModel):
    role: str
    institution_name: str

class SubmitAttemptRequest(BaseModel):
    student_name: str
    answers: Dict[str, str] 
    challenge_code: Optional[str] = None

class ChallengeCreateRequest(BaseModel):
    quiz_id: int

class FlashcardReviewRequest(BaseModel):
    flashcard_id: int
    quality: int

class ExportQuizRequest(BaseModel):
    include_answer_key: bool = False

class CreateClassroomRequest(BaseModel):
    name: str

class JoinClassroomRequest(BaseModel):
    join_code: str

class BrandingRequest(BaseModel):
    academy_name: str
    logo_path: str = ""

# 🪄 PHASE 1 NEW SCHEMAS (For Editing & Smart Questions)
class UpdateQuizRequest(BaseModel):
    quiz_data: Dict[str, Any]
    exam_metadata: Dict[str, Any]

class RegenerateQuestionRequest(BaseModel):
    question_type: str # 'mcq', 'fill_blank', 'short_answer', 'long_answer'
    difficulty: str
    question_style: str = "Auto"

class BookmarkRequest(BaseModel):
    quiz_id: int
    question_type: str
    question_data: Dict[str, Any]

# --- AUTH DEPENDENCIES ---
def get_current_user(authorization: Optional[str] = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header.")
    token = authorization.removeprefix("Bearer ").strip()
    user_id = database.get_user_id_for_token(token)
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid or expired session token.")
    user = database.get_user_by_id(user_id)
    if not user:
        raise HTTPException(status_code=401, detail="User not found.")
    return user

def get_optional_user(authorization: Optional[str] = Header(None)):
    if not authorization or not authorization.startswith("Bearer "):
        return None
    token = authorization.removeprefix("Bearer ").strip()
    user_id = database.get_user_id_for_token(token)
    if user_id:
        return database.get_user_by_id(user_id)
    return None

def require_teacher(user=Depends(get_current_user)):
    if user["role"] not in ("teacher", "admin"):
        raise HTTPException(status_code=403, detail="Only teachers/admins can perform this action.")
    return user

# --- AUTH ENDPOINTS ---
@app.post("/auth/register")
def register(req: RegisterRequest):
    if database.get_user_by_email(req.email):
        raise HTTPException(status_code=409, detail="An account with this email already exists.")
    password_hash, salt = auth.hash_password(req.password)
    user_id = database.create_user(req.name, req.email, password_hash, salt, "unassigned", "")
    token = auth.generate_token()
    database.create_session(token, user_id)
    return {"token": token, "user_id": user_id, "role": "unassigned", "name": req.name}

@app.post("/auth/login")
def login(req: LoginRequest):
    user = database.get_user_by_email(req.email)
    if not user or not auth.verify_password(req.password, user["password_hash"], user["salt"]):
        raise HTTPException(status_code=401, detail="Invalid email or password.")
    token = auth.generate_token()
    database.create_session(token, user["id"])
    return {"token": token, "user_id": user["id"], "role": user["role"], "name": user["name"]}

@app.post("/auth/update-profile")
def update_profile(req: UpdateProfileRequest, user=Depends(get_current_user)):
    database.update_user_profile(user["id"], req.role, req.institution_name)
    return {"message": "Profile updated successfully", "role": req.role}

@app.post("/auth/logout")
def logout(authorization: Optional[str] = Header(None)):
    if authorization and authorization.startswith("Bearer "):
        auth_token = authorization.removeprefix("Bearer ").strip()
        database.delete_session(auth_token)
    return {"ok": True}

@app.post("/quiz/analyze-document")
async def analyze_document(file: UploadFile = File(...)):
    try:
        file_bytes = await file.read()
        total_pages = get_document_page_count(file_bytes, file.filename)
        return {"total_pages": total_pages}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

# ==========================================
# 📸 MULTI-SOURCE QUIZ INTEGRATED ENDPOINT (UPDATED)
# ==========================================
@app.post("/quiz/generate")
async def generate_quiz(
    request: Request,
    file: Optional[UploadFile] = File(None), 
    files: Optional[List[UploadFile]] = File(None), 
    file_ranges: Optional[str] = Form(None), 
    youtube_url: str = Form(""),
    yt_start_min: int = Form(0),    
    yt_end_min: int = Form(0),      
    num_mcq: int = Form(0),
    num_fill_blank: int = Form(0),
    num_short: int = Form(0),
    num_long: int = Form(0),
    difficulty: str = Form("Medium"),
    question_style: str = Form("Auto"),  # 🧠 NEW: Smart Question Style parameter
    start_page: int = Form(1),      
    end_page: int = Form(1000),     
    institution_name: str = Form(""),
    department: str = Form(""),
    subject: str = Form(""),
    class_name: str = Form(""),
    teacher_name: str = Form(""),
    exam_title: str = Form(""),
    user=Depends(get_current_user), 
):
    req_id = getattr(request.state, "req_id", "Unknown")
    
    if num_mcq + num_fill_blank + num_short + num_long <= 0:
        raise HTTPException(status_code=400, detail="Request at least 1 question.")

    raw_text = ""
    source_identifiers = []

    all_uploaded_files = []
    if file:
        all_uploaded_files.append(file)
    if files:
        all_uploaded_files.extend(files)

    parsed_ranges = {}
    if file_ranges:
        try:
            ranges_list = json.loads(file_ranges)
            for r in ranges_list:
                parsed_ranges[r.get("filename", "")] = r
        except Exception as e:
            logger.error(f"Failed to parse file_ranges JSON: {e}")

    if youtube_url.strip():
        source_identifiers.append(f"{youtube_url.strip()}_{yt_start_min}_{yt_end_min}")
        raw_text = get_youtube_transcript(youtube_url.strip(), start_min=yt_start_min, end_min=yt_end_min)
        
    elif all_uploaded_files:
        for f in all_uploaded_files:
            if not f.filename: continue
            
            f_bytes = await f.read()
            f_start = parsed_ranges.get(f.filename, {}).get("start", start_page)
            f_end = parsed_ranges.get(f.filename, {}).get("end", end_page)

            if f.content_type.startswith('image/'):
                try:
                    image = Image.open(io.BytesIO(f_bytes))
                    if image.mode != 'RGB': image = image.convert('RGB')
                    extracted = pytesseract.image_to_string(image)
                    if len(extracted.strip()) < 50:
                        raise HTTPException(status_code=422, detail=f"Text unreadable in image: {f.filename}")
                    raw_text += extracted + "\n\n"
                    source_identifiers.append(hashlib.md5(f_bytes).hexdigest())
                except Exception as e:
                    raise HTTPException(status_code=422, detail=f"Image processing failed for {f.filename}")
            else:
                try:
                    extracted = extract_text_from_document(f_bytes, f.filename, start_page=f_start, end_page=f_end)
                    if extracted.startswith("Error"):
                        raise HTTPException(status_code=422, detail=extracted)
                    raw_text += extracted + "\n\n"
                    source_identifiers.append(hashlib.md5(f_bytes).hexdigest())
                except Exception as e:
                    raise HTTPException(status_code=422, detail=f"Failed to process {f.filename}")
    else:
        raise HTTPException(status_code=400, detail="Provide either Document(s), Image(s), or a YouTube URL.")

    if not raw_text.strip() or raw_text.startswith("Error"):
        raise HTTPException(status_code=422, detail="Combined text extraction failed. Content might be too short.")

    question_counts = {
        "mcq": num_mcq, "fill_blank": num_fill_blank, "short_answer": num_short, "long_answer": num_long
    }

    combined_identifier = "_".join(source_identifiers)
    req_hash = generate_hash(
        source_identifier=combined_identifier, mcq=num_mcq, fill_blank=num_fill_blank,
        short_ans=num_short, long_ans=num_long, difficulty=difficulty
    )

    quiz_data = get_cached_quiz(req_hash)
    if not quiz_data:
        # 🧠 Pass the question_style down to the generator
        quiz_data = generate_quiz_from_large_text(raw_text, question_counts, difficulty, question_style=question_style)
        save_quiz_to_cache(req_hash, quiz_data)

    calculated_marks = (num_mcq * 1) + (num_fill_blank * 1) + (num_short * 2) + (num_long * 5)
    calculated_duration = (num_mcq * 1) + (num_fill_blank * 1) + (num_short * 3) + (num_long * 8)
    if calculated_duration < 15: calculated_duration = 15

    clean_title = exam_title.strip() if exam_title.strip() else "Assessment Examination"

    exam_metadata = {
        "institution_name": institution_name or "Academy Worksheet",
        "department": department, 
        "subject": subject, 
        "class_name": class_name,
        "teacher_name": teacher_name or user["name"], 
        "exam_title": clean_title,
        "duration_minutes": calculated_duration, 
        "total_marks": calculated_marks,
        "source_text_context": raw_text[:20000] # 🧠 Save context for Phase 1 Regeneration securely
    }

    quiz_id = database.create_quiz(user["id"], exam_metadata, quiz_data)
    return {"quiz_id": quiz_id, "exam_metadata": exam_metadata, "quiz_data": quiz_data}


# ==========================================
# 🪄 PHASE 1 NEW APIS: EDITOR & REGENERATION
# ==========================================
@app.put("/quiz/{quiz_id}")
def update_quiz_data(quiz_id: int, req: UpdateQuizRequest, user=Depends(get_current_user)):
    """Updates edited quiz data manually (Reorder, Edit, Add, Delete) directly in DB."""
    conn = database.get_db_connection()
    cursor = conn.cursor()
    # Check ownership
    cursor.execute("SELECT id FROM quizzes WHERE id = ? AND teacher_id = ?", (quiz_id, user["id"]))
    if not cursor.fetchone(): 
        conn.close()
        raise HTTPException(status_code=403, detail="Unauthorized or Quiz not found.")
    
    cursor.execute("UPDATE quizzes SET quiz_data = ?, exam_metadata = ? WHERE id = ?", 
                  (json.dumps(req.quiz_data), json.dumps(req.exam_metadata), quiz_id))
    conn.commit()
    conn.close()
    return {"message": "Quiz updated successfully."}

@app.post("/quiz/{quiz_id}/regenerate-question")
def regenerate_single_question(quiz_id: int, req: RegenerateQuestionRequest, user=Depends(get_current_user)):
    """Re-uses existing AI pipeline to fetch exactly ONE new question of requested type."""
    quiz = database.get_quiz(quiz_id)
    if not quiz or quiz["teacher_id"] != user["id"]: 
        raise HTTPException(status_code=403, detail="Unauthorized.")
    
    context = quiz["exam_metadata"].get("source_text_context", "")
    if not context: 
        raise HTTPException(status_code=400, detail="Original context not found. Cannot regenerate.")

    q_counts = {"mcq": 0, "fill_blank": 0, "short_answer": 0, "long_answer": 0}
    if req.question_type in q_counts:
        q_counts[req.question_type] = 1
    
    # EXISTING PIPELINE CALL (No new prompts required!)
    new_data = generate_quiz_from_large_text(context, q_counts, req.difficulty, question_style=req.question_style)
    
    key_map = {"mcq": "mcq_questions", "fill_blank": "fill_blank_questions", "short_answer": "short_questions", "long_answer": "long_questions"}
    target_key = key_map.get(req.question_type)
    
    if new_data and target_key and new_data.get(target_key) and len(new_data[target_key]) > 0:
        return {"new_question": new_data[target_key][0]}
    else:
        raise HTTPException(status_code=500, detail="AI Engine failed to regenerate question.")

@app.post("/bookmarks")
def save_bookmark(req: BookmarkRequest, user=Depends(get_current_user)):
    """Saves a specific question to the user's bookmarks."""
    conn = database.get_db_connection()
    cursor = conn.cursor()
    cursor.execute("INSERT INTO bookmarked_questions (user_id, quiz_id, question_type, question_data) VALUES (?, ?, ?, ?)",
                   (user["id"], req.quiz_id, req.question_type, json.dumps(req.question_data)))
    conn.commit()
    conn.close()
    return {"message": "Question bookmarked successfully!"}

@app.get("/bookmarks")
def get_bookmarks(user=Depends(get_current_user)):
    """Fetches all saved questions."""
    conn = database.get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM bookmarked_questions WHERE user_id = ? ORDER BY created_at DESC", (user["id"],))
    bookmarks = []
    for row in cursor.fetchall():
        bookmarks.append({
            "id": row["id"], "quiz_id": row["quiz_id"], 
            "question_type": row["question_type"], "question_data": json.loads(row["question_data"]),
            "date": row["created_at"]
        })
    conn.close()
    return {"bookmarks": bookmarks}


# ==========================================
# 🖨️ TEACHER EXPORT ENGINE (DOCX & PDF VIP LAYOUT)
# ==========================================
def process_base64_logo_safe(branding):
    """Safely converts Base64 logo to a temporary physical file for bulletproof PDF/Word embedding."""
    if not branding or not branding.get("logo_path") or not branding["logo_path"].startswith("data:image"):
        return None
    try:
        b64_data = branding["logo_path"].split(",")[1]
        img_data = base64.b64decode(b64_data)
        pil_img = Image.open(io.BytesIO(img_data))
        
        if pil_img.mode in ("RGBA", "P", "LA"):
            background = Image.new("RGB", pil_img.size, (255, 255, 255))
            if 'A' in pil_img.getbands():
                background.paste(pil_img, mask=pil_img.split()[-1])
            else:
                background.paste(pil_img)
            pil_img = background
        elif pil_img.mode != "RGB":
            pil_img = pil_img.convert("RGB")
            
        fd, path = tempfile.mkstemp(suffix=".jpg")
        with os.fdopen(fd, 'wb') as f:
            pil_img.save(f, format="JPEG", quality=95)
        return path
    except Exception as e:
        logger.error(f"Logo processing error: {e}")
        return None

@app.post("/quiz/{quiz_id}/export/docx")
def export_quiz_docx(quiz_id: int, req: ExportQuizRequest, user=Depends(require_teacher)):
    quiz = database.get_quiz(quiz_id)
    if not quiz: raise HTTPException(status_code=404, detail="Quiz not found.")
    
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    meta = quiz["exam_metadata"]
    quiz_data = quiz["quiz_data"]
    
    branding = database.get_teacher_branding(user["id"])
    academy_name = branding["academy_name"] if branding and branding.get("academy_name") else meta.get("institution_name", "Academy Worksheet")
    
    logo_path = process_base64_logo_safe(branding)

    header_table = doc.add_table(rows=1, cols=2)
    header_table.columns[0].width = Inches(1.2)
    header_table.columns[1].width = Inches(5.5)

    if logo_path and os.path.exists(logo_path):
        try:
            p_logo = header_table.cell(0, 0).paragraphs[0]
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Inches(1.0))
        except Exception as e:
            logger.error(f"Failed to add logo to DOCX: {e}")

    p_title = header_table.cell(0, 1).paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_uni = p_title.add_run(f"{academy_name}\n")
    r_uni.bold = True
    r_uni.font.size = Pt(16)
    
    exam_title = meta.get("exam_title", "Assessment Examination")
    r_exam = p_title.add_run(f"{exam_title}\n\n")
    r_exam.bold = True
    r_exam.font.size = Pt(13)
    
    t_name = str(meta.get('teacher_name', user['name'])).upper()
    subject_val = meta.get('subject', 'Not Specified')
    class_val = meta.get('class_name', 'Not Specified')
    dur = meta.get('duration_minutes', 60)
    marks = meta.get('total_marks', 100)
    
    r_meta1 = p_title.add_run(f"Subject: {subject_val}   |   Class/Semester: {class_val}\n")
    r_meta1.font.size = Pt(10)
    
    r_meta2 = p_title.add_run(f"Teacher: ")
    r_meta2.font.size = Pt(10)
    r_meta_tname = p_title.add_run(f"{t_name}")
    r_meta_tname.bold = True
    try:
        r_meta_tname.font.color.rgb = RGBColor(0, 0, 139) 
    except:
        pass
    
    r_meta3 = p_title.add_run(f"   |   Duration: {dur} mins   |   Total Marks: {marks}")
    r_meta3.font.size = Pt(10)

    doc.add_paragraph("-" * 80)
    doc.add_paragraph("Student Name: __________________________________   Roll No: ____________   Date: ___________")
    doc.add_paragraph()
    
    if quiz_data.get("mcq_questions"):
        doc.add_heading("Section A: Multiple Choice Questions", level=2)
        counter = 1
        for q in quiz_data["mcq_questions"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run(f"Q{counter}. ").bold = True
            p.add_run(q['question_text'])
            for j, opt in enumerate(q['options']):
                p_opt = doc.add_paragraph(f"   {chr(65+j)}. {opt}")
                p_opt.paragraph_format.space_after = Pt(2)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            counter += 1

    if quiz_data.get("fill_blank_questions"):
        doc.add_heading("Section B: Fill in the Blanks", level=2)
        counter = 1
        for q in quiz_data["fill_blank_questions"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.add_run(f"Q{counter}. ").bold = True
            p.add_run(q['question_text'])
            counter += 1

    if quiz_data.get("short_questions"):
        doc.add_heading("Section C: Short Answer Questions", level=2)
        counter = 1
        for q in quiz_data["short_questions"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(24) 
            p.add_run(f"Q{counter}. ").bold = True
            p.add_run(q['question_text'])
            counter += 1

    if quiz_data.get("long_questions"):
        doc.add_heading("Section D: Detailed Explanation", level=2)
        counter = 1
        for q in quiz_data["long_questions"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(48) 
            p.add_run(f"Q{counter}. ").bold = True
            p.add_run(q['question_text'])
            counter += 1

    if req.include_answer_key:
        doc.add_page_break()
        doc.add_heading("Answer Key", level=1)
        for sec in ["mcq_questions", "fill_blank_questions", "short_questions", "long_questions"]:
            if quiz_data.get(sec):
                k = 1
                for q in quiz_data[sec]:
                    ans = q.get('correct_answer') or q.get('model_answer')
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.add_run(f"Q{k}. ").bold = True
                    p.add_run(str(ans))
                    k += 1

    if logo_path and os.path.exists(logo_path):
        os.remove(logo_path)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=Quiz_{quiz_id}.docx"})

@app.post("/quiz/{quiz_id}/export/pdf")
def export_quiz_pdf(quiz_id: int, req: ExportQuizRequest, user=Depends(require_teacher)):
    quiz = database.get_quiz(quiz_id)
    if not quiz: raise HTTPException(status_code=404, detail="Quiz not found.")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    
    q_style = ParagraphStyle(name='Question', parent=styles['Normal'], fontName='Helvetica', spaceBefore=8, spaceAfter=2, leading=14)
    opt_style = ParagraphStyle(name='Option', parent=styles['Normal'], leftIndent=20, spaceAfter=2, leading=12)
    
    elements = []
    meta = quiz["exam_metadata"]
    quiz_data = quiz["quiz_data"]
    
    branding = database.get_teacher_branding(user["id"])
    academy_name = branding["academy_name"] if branding and branding.get("academy_name") else meta.get("institution_name", "Academy Worksheet")
    
    logo_path = process_base64_logo_safe(branding)
    logo_img = None

    if logo_path and os.path.exists(logo_path):
        try:
            logo_img = RLImage(logo_path, width=1.0*72, height=1.0*72, kind='proportional')
        except Exception as e:
            logger.error(f"Failed to add logo to PDF: {e}")

    exam_title = meta.get("exam_title", "Assessment Examination")
    t_name = str(meta.get('teacher_name', user['name'])).upper()
    subject_val = meta.get('subject', 'Not Specified')
    class_val = meta.get('class_name', 'Not Specified')
    dur = meta.get('duration_minutes', 60)
    marks = meta.get('total_marks', 100)
    
    p_header = Paragraph(f"""
        <font size="16"><b>{academy_name}</b></font><br/>
        <font size="13"><b>{exam_title}</b></font><br/><br/>
        <font size="10">Subject: {subject_val} &nbsp;&nbsp;|&nbsp;&nbsp; Class/Semester: {class_val}</font><br/>
        <font size="10">Teacher: <b><font color='blue'>{t_name}</font></b> &nbsp;&nbsp;|&nbsp;&nbsp; Duration: {dur} mins &nbsp;&nbsp;|&nbsp;&nbsp; Total Marks: {marks}</font>
    """, ParagraphStyle(name='HeaderStyle', alignment=TA_CENTER, leading=14))

    if logo_img:
        header_table = Table([[logo_img, p_header]], colWidths=[1.2*72, 6.0*72])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('ALIGN', (0,0), (0,0), 'CENTER'),
            ('ALIGN', (1,0), (1,0), 'CENTER')
        ]))
    else:
        header_table = Table([[p_header]], colWidths=[7.2*72])
        header_table.setStyle(TableStyle([('ALIGN', (0,0), (-1,-1), 'CENTER')]))
    
    elements.append(header_table)
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("_" * 70, ParagraphStyle(name='Line', alignment=TA_CENTER)))
    elements.append(Spacer(1, 5))
    elements.append(Paragraph("<b>Student Name:</b> ___________________________   <b>Roll No:</b> _____________   <b>Date:</b> ___________", styles['Normal']))
    elements.append(Spacer(1, 15))
    
    if quiz_data.get("mcq_questions"):
        elements.append(Paragraph("<b>Section A: Multiple Choice</b>", styles['Heading3']))
        counter = 1
        for q in quiz_data["mcq_questions"]:
            elements.append(Paragraph(f"<b>Q{counter}.</b> {q['question_text']}", q_style))
            for j, opt in enumerate(q['options']):
                elements.append(Paragraph(f"{chr(65+j)}. {opt}", opt_style))
            counter += 1
            elements.append(Spacer(1, 4))

    if quiz_data.get("fill_blank_questions"):
        elements.append(Paragraph("<b>Section B: Fill in the Blanks</b>", styles['Heading3']))
        counter = 1
        for q in quiz_data["fill_blank_questions"]:
            elements.append(Paragraph(f"<b>Q{counter}.</b> {q['question_text']}", q_style))
            elements.append(Spacer(1, 12)) 
            counter += 1

    if quiz_data.get("short_questions"):
        elements.append(Paragraph("<b>Section C: Short Answer</b>", styles['Heading3']))
        counter = 1
        for q in quiz_data["short_questions"]:
            elements.append(Paragraph(f"<b>Q{counter}.</b> {q['question_text']}", q_style))
            elements.append(Spacer(1, 30)) 
            counter += 1

    if quiz_data.get("long_questions"):
        elements.append(Paragraph("<b>Section D: Detailed Explanation</b>", styles['Heading3']))
        counter = 1
        for q in quiz_data["long_questions"]:
            elements.append(Paragraph(f"<b>Q{counter}.</b> {q['question_text']}", q_style))
            elements.append(Spacer(1, 50)) 
            counter += 1

    if req.include_answer_key:
        elements.append(PageBreak())
        elements.append(Paragraph("Answer Key", styles['Heading2']))
        for sec in ["mcq_questions", "fill_blank_questions", "short_questions", "long_questions"]:
            if quiz_data.get(sec):
                k = 1
                for q in quiz_data[sec]:
                    ans = q.get('correct_answer') or q.get('model_answer')
                    elements.append(Paragraph(f"<b>Q{k}.</b> {ans}", styles['Normal']))
                    elements.append(Spacer(1, 4))
                    k += 1

    doc.build(elements)
    buffer.seek(0)

    if logo_path and os.path.exists(logo_path):
        os.remove(logo_path)

    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=Quiz_{quiz_id}.pdf"})

# ==========================================
# Baqi Code (Teacher Quizzes, Classes, Submissions, Analytics, Branding waghaira)
# ==========================================
@app.get("/quiz/{quiz_id}")
def get_quiz_for_student(quiz_id: int, request: Request):
    quiz = database.get_quiz(quiz_id)
    if not quiz: raise HTTPException(status_code=404, detail="Quiz not found.")
    quiz_data = quiz["quiz_data"]
    safe_quiz_data = {
        "mcq_questions": [{"question_text": q["question_text"], "options": q["options"]} for q in quiz_data.get("mcq_questions", [])],
        "fill_blank_questions": [{"question_text": q["question_text"]} for q in quiz_data.get("fill_blank_questions", [])],
        "short_questions": [{"question_text": q["question_text"]} for q in quiz_data.get("short_questions", [])],
        "long_questions": [{"question_text": q["question_text"]} for q in quiz_data.get("long_questions", [])],
    }
    return {"quiz_id": quiz_id, "exam_metadata": quiz["exam_metadata"], "quiz_data": safe_quiz_data}

@app.post("/quiz/{quiz_id}/submit")
def submit_attempt(quiz_id: int, req: SubmitAttemptRequest, request: Request, user=Depends(get_optional_user)):
    quiz = database.get_quiz(quiz_id)
    if not quiz: raise HTTPException(status_code=404, detail="Quiz not found.")
    quiz_data = quiz["quiz_data"]
    answers = req.answers
    results = {"mcq": [], "fill_blank": [], "short": [], "long": [], "total_score": 0.0, "max_score": 0}

    for i, q in enumerate(quiz_data.get("mcq_questions", [])):
        selected = answers.get(f"mcq_{i}")
        is_correct = check_mcq(selected, q["correct_answer"])
        results["mcq"].append({"question": q["question_text"], "selected": selected, "correct_answer": q["correct_answer"], "is_correct": is_correct, "explanation": q["explanation"]})
        results["max_score"] += 1
        results["total_score"] += 1 if is_correct else 0

    for i, q in enumerate(quiz_data.get("fill_blank_questions", [])):
        student_ans = answers.get(f"fb_{i}", "")
        is_correct = check_fill_blank(student_ans, q["correct_answer"])
        results["fill_blank"].append({"question": q["question_text"], "student_answer": student_ans, "correct_answer": q["correct_answer"], "is_correct": is_correct, "explanation": q["explanation"]})
        results["max_score"] += 1
        results["total_score"] += 1 if is_correct else 0

    for i, q in enumerate(quiz_data.get("short_questions", [])):
        student_ans = answers.get(f"short_{i}", "")
        grade = grade_long_answer(q["question_text"], q["correct_answer"], [], student_ans)
        results["short"].append({"question": q["question_text"], "student_answer": student_ans, "model_answer": q["correct_answer"], "score_percent": grade["score_percent"], "feedback": grade["feedback"]})
        results["max_score"] += 2
        results["total_score"] += (grade["score_percent"] / 100) * 2

    for i, q in enumerate(quiz_data.get("long_questions", [])):
        student_ans = answers.get(f"long_{i}", "")
        grade = grade_long_answer(q["question_text"], q["model_answer"], q.get("key_points", []), student_ans)
        results["long"].append({"question": q["question_text"], "student_answer": student_ans, "model_answer": q["model_answer"], "score_percent": grade["score_percent"], "feedback": grade["feedback"]})
        results["max_score"] += 5
        results["total_score"] += (grade["score_percent"] / 100) * 5

    attempt_id = database.save_attempt(quiz_id, req.student_name, answers, results)
    
    if user: database.update_streak_and_badges(user["id"])
    if user and req.challenge_code:
        challenge = database.get_challenge_by_code(req.challenge_code)
        if challenge: database.record_challenge_participant(challenge["id"], user["id"], attempt_id)

    return {"attempt_id": attempt_id, "results": results}

@app.post("/challenge/create")
def create_challenge(req: ChallengeCreateRequest, user=Depends(get_current_user)):
    quiz = database.get_quiz(req.quiz_id)
    if not quiz: raise HTTPException(status_code=404, detail="Quiz not found.")
    code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
    challenge_id = database.create_challenge(user["id"], req.quiz_id, code)
    return {"challenge_id": challenge_id, "code": code, "share_text": f"Join using code: {code}"}

@app.get("/challenge/{code}")
def get_challenge(code: str, request: Request):
    challenge = database.get_challenge_by_code(code.upper())
    if not challenge: raise HTTPException(status_code=404, detail="Invalid Challenge Code.")
    return get_quiz_for_student(challenge["quiz_id"], request)

@app.get("/challenge/{code}/leaderboard")
def get_challenge_leaderboard(code: str):
    challenge = database.get_challenge_by_code(code.upper())
    if not challenge: raise HTTPException(status_code=404, detail="Invalid Challenge.")
    return {"code": code, "leaderboard": database.get_challenge_leaderboard(challenge["id"])}

@app.get("/user/dashboard")
def get_student_dashboard(user=Depends(get_current_user)):
    return database.get_user_gamification(user["id"])

@app.post("/quiz/{quiz_id}/flashcards")
def create_flashcards(quiz_id: int, user=Depends(get_current_user)):
    quiz = database.get_quiz(quiz_id)
    if not quiz: raise HTTPException(status_code=404, detail="Quiz not found.")
    for q in quiz.get("quiz_data", {}).get("mcq_questions", []):
        database.create_flashcard(quiz_id, user["id"], front=q["question_text"], back=q["correct_answer"], q_type="mcq")
    for q in quiz.get("quiz_data", {}).get("short_questions", []):
        database.create_flashcard(quiz_id, user["id"], front=q["question_text"], back=q["correct_answer"], q_type="short")
    return {"message": "Flashcards generated."}

@app.get("/flashcards/due")
def get_due_flashcards(user=Depends(get_current_user)):
    due_cards = database.get_due_flashcards(user["id"])
    return {"due_count": len(due_cards), "flashcards": due_cards}

@app.post("/flashcards/review")
def review_flashcard(req: FlashcardReviewRequest, user=Depends(get_current_user)):
    database.update_flashcard_sm2(user["id"], req.flashcard_id, req.quality)
    return {"message": "Review recorded."}

@app.get("/teacher/quizzes")
def get_teacher_quizzes_api(user=Depends(require_teacher)):
    quizzes = database.get_quizzes_for_teacher(user["id"])
    formatted_quizzes = []
    for q in quizzes:
        formatted_quizzes.append({
            "id": q["id"],
            "title": q["exam_metadata"].get("exam_title", "Assessment Examination"),
            "subject": q["exam_metadata"].get("subject", "Not Specified"),
            "date": q["created_at"].split(" ")[0] 
        })
    return {"quizzes": formatted_quizzes}

@app.delete("/quiz/{quiz_id}")
def delete_quiz_api(quiz_id: int, user=Depends(require_teacher)):
    conn = database.get_db_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM quizzes WHERE id = ? AND teacher_id = ?", (quiz_id, user["id"]))
    conn.commit()
    deleted = cursor.rowcount > 0
    conn.close()
    if not deleted: raise HTTPException(status_code=404, detail="Quiz not found or unauthorized.")
    return {"message": "Quiz deleted successfully"}

@app.get("/teacher/overview")
def get_teacher_overview(user=Depends(require_teacher)):
    conn = database.get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(id) as count FROM quizzes WHERE teacher_id = ?", (user["id"],))
    total_quizzes = cursor.fetchone()["count"]
    cursor.execute("SELECT COUNT(id) as count FROM classrooms WHERE teacher_id = ?", (user["id"],))
    total_classes = cursor.fetchone()["count"]
    cursor.execute('''
        SELECT COUNT(a.id) as attempt_count, 
               AVG(CAST(json_extract(a.results, '$.total_score') AS REAL) / 
                   CAST(json_extract(a.results, '$.max_score') AS REAL)) * 100 as avg_score
        FROM attempts a
        JOIN quizzes q ON a.quiz_id = q.id
        WHERE q.teacher_id = ? AND json_extract(a.results, '$.max_score') > 0
    ''', (user["id"],))
    stats = cursor.fetchone()
    total_attempts = stats["attempt_count"] if stats["attempt_count"] else 0
    avg_score = round(stats["avg_score"] or 0, 1)
    conn.close()
    return {
        "total_quizzes": total_quizzes,
        "total_classes": total_classes,
        "total_attempts": total_attempts,
        "avg_score": avg_score
    }

@app.post("/teacher/classrooms")
def create_classroom_api(req: CreateClassroomRequest, user=Depends(require_teacher)):
    join_code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=7))
    class_id = database.create_classroom(user["id"], req.name, join_code)
    return {
        "message": "Classroom created successfully", 
        "class_id": class_id, 
        "join_code": join_code,
        "name": req.name
    }

@app.get("/teacher/classrooms")
def get_classrooms_api(user=Depends(require_teacher)):
    conn = database.get_db_connection()
    cursor = conn.cursor()
    cursor.execute('''
        SELECT c.id, c.name, c.join_code, c.created_at, 
               (SELECT COUNT(id) FROM classroom_students WHERE classroom_id = c.id) as student_count
        FROM classrooms c
        WHERE c.teacher_id = ?
        ORDER BY c.created_at DESC
    ''', (user["id"],))
    classes = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {"classes": classes}

@app.post("/student/classrooms/join")
def join_classroom_api(req: JoinClassroomRequest, user=Depends(get_current_user)):
    if user["role"] == "teacher":
        raise HTTPException(status_code=400, detail="Teachers cannot join classes as students.")
    result = database.join_classroom(user["id"], req.join_code.upper())
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return {"message": "Successfully joined the classroom!", "classroom_id": result["classroom_id"]}

@app.get("/teacher/analytics/recent-attempts")
def get_recent_attempts_api(user=Depends(require_teacher)):
    conn = database.get_db_connection()
    cursor = conn.cursor()
    cursor.execute('''
        SELECT a.id, a.student_name, a.created_at, q.exam_metadata, a.results
        FROM attempts a
        JOIN quizzes q ON a.quiz_id = q.id
        WHERE q.teacher_id = ?
        ORDER BY a.id DESC
        LIMIT 15
    ''', (user["id"],))
    attempts = []
    for row in cursor.fetchall():
        meta = json.loads(row["exam_metadata"])
        res = json.loads(row["results"])
        max_score = res.get("max_score", 0)
        score = res.get("total_score", 0)
        score_pct = round((score / max_score) * 100, 1) if max_score > 0 else 0
        attempts.append({
            "id": row["id"],
            "student_name": row["student_name"],
            "quiz_title": meta.get("exam_title", "Untitled Quiz"),
            "score_percent": score_pct,
            "date": row["created_at"].split(" ")[0]
        })
    conn.close()
    return {"attempts": attempts}

@app.get("/teacher/branding")
def get_branding_api(user=Depends(require_teacher)):
    branding = database.get_teacher_branding(user["id"])
    return branding or {"academy_name": "", "logo_path": ""}

@app.post("/teacher/branding")
def update_branding_api(req: BrandingRequest, user=Depends(require_teacher)):
    database.update_teacher_branding(user["id"], req.academy_name, req.logo_path)
    return {"message": "Academy branding updated successfully!"}