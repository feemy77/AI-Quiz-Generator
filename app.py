import streamlit as st
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pdf_extractor import extract_text_from_pdf
from quiz_generator import generate_quiz_from_large_text
from grading_engine import check_mcq, check_fill_blank, grade_long_answer

PLACEHOLDER = "-- Select an option --"

# ==========================================
# 1. Helper Function: Create Word Document
# ==========================================
def create_word_docx(quiz_data, exam_metadata=None):
    doc = Document()
    exam_metadata = exam_metadata or {}

    # ---------- Header / Branding ----------
    if exam_metadata.get("institution_name"):
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = h.add_run(exam_metadata["institution_name"])
        run.bold = True
        run.font.size = doc.styles['Heading 1'].font.size

    if exam_metadata.get("department"):
        p = doc.add_paragraph(exam_metadata["department"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = exam_metadata.get("exam_title") or "AI Generated Quiz"
    doc.add_heading(title, level=0 if not exam_metadata.get("institution_name") else 1)

    # Info line: Subject | Class | Teacher | Duration | Total Marks
    info_bits = []
    if exam_metadata.get("subject"):
        info_bits.append(f"Subject: {exam_metadata['subject']}")
    if exam_metadata.get("class_name"):
        info_bits.append(f"Class: {exam_metadata['class_name']}")
    if exam_metadata.get("teacher_name"):
        info_bits.append(f"Examiner: {exam_metadata['teacher_name']}")
    if exam_metadata.get("duration_minutes"):
        info_bits.append(f"Duration: {exam_metadata['duration_minutes']} mins")
    if exam_metadata.get("total_marks"):
        info_bits.append(f"Total Marks: {exam_metadata['total_marks']}")
    if info_bits:
        doc.add_paragraph(" | ".join(info_bits))

    doc.add_paragraph("Student Name: ____________________     Roll No: ____________________")
    doc.add_paragraph("")

    mcqs = quiz_data.get("mcq_questions", [])
    fill_blanks = quiz_data.get("fill_blank_questions", [])
    longs = quiz_data.get("long_questions", [])

    # ---------- Section: MCQs ----------
    if mcqs:
        doc.add_heading('Section A: Multiple Choice Questions', level=1)
        for i, q in enumerate(mcqs):
            doc.add_paragraph(f"Q{i+1}: {q['question_text']}")
            for j, opt in enumerate(q["options"]):
                letter = chr(65 + j)
                doc.add_paragraph(f"    {letter}) {opt}")
            doc.add_paragraph("")

    # ---------- Section: Fill in the Blanks ----------
    if fill_blanks:
        doc.add_heading('Section B: Fill in the Blanks', level=1)
        for i, q in enumerate(fill_blanks):
            doc.add_paragraph(f"Q{i+1}: {q['question_text']}")
        doc.add_paragraph("")

    # ---------- Section: Long Questions ----------
    if longs:
        doc.add_heading('Section C: Long Questions', level=1)
        for i, q in enumerate(longs):
            doc.add_paragraph(f"Q{i+1}: {q['question_text']}")
        doc.add_paragraph("")

    doc.add_page_break()

    # ---------- Answer Key ----------
    doc.add_heading('Answer Key & Explanations', level=1)

    if mcqs:
        doc.add_heading('Section A: MCQ Answers', level=2)
        for i, q in enumerate(mcqs):
            p = doc.add_paragraph()
            p.add_run(f"Q{i+1} Answer: ").bold = True
            p.add_run(f"{q['correct_answer']}\n")
            p.add_run("Explanation: ").italic = True
            p.add_run(f"{q['explanation']}")
            doc.add_paragraph("")

    if fill_blanks:
        doc.add_heading('Section B: Fill in the Blank Answers', level=2)
        for i, q in enumerate(fill_blanks):
            p = doc.add_paragraph()
            p.add_run(f"Q{i+1} Answer: ").bold = True
            p.add_run(f"{q['correct_answer']}\n")
            p.add_run("Explanation: ").italic = True
            p.add_run(f"{q['explanation']}")
            doc.add_paragraph("")

    if longs:
        doc.add_heading('Section C: Long Question Model Answers', level=2)
        for i, q in enumerate(longs):
            p = doc.add_paragraph()
            p.add_run(f"Q{i+1} Model Answer: ").bold = True
            p.add_run(f"{q['model_answer']}\n")
            p.add_run("Key Points: ").italic = True
            p.add_run(", ".join(q.get("key_points", [])))
            doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==========================================
# 2. Helper Function: Grade a submitted test
# ==========================================
def grade_test(quiz_data, answers):
    mcqs = quiz_data.get("mcq_questions", [])
    fill_blanks = quiz_data.get("fill_blank_questions", [])
    longs = quiz_data.get("long_questions", [])

    results = {"mcq": [], "fill_blank": [], "long": [], "total_score": 0.0, "max_score": 0}

    for i, q in enumerate(mcqs):
        selected = answers.get(f"mcq_{i}")
        if selected == PLACEHOLDER:
            selected = None
        is_correct = check_mcq(selected, q["correct_answer"])
        results["mcq"].append({
            "question": q["question_text"],
            "selected": selected,
            "correct_answer": q["correct_answer"],
            "is_correct": is_correct,
            "explanation": q["explanation"],
        })
        results["max_score"] += 1
        results["total_score"] += 1 if is_correct else 0

    for i, q in enumerate(fill_blanks):
        student_ans = answers.get(f"fb_{i}", "")
        is_correct = check_fill_blank(student_ans, q["correct_answer"])
        results["fill_blank"].append({
            "question": q["question_text"],
            "student_answer": student_ans,
            "correct_answer": q["correct_answer"],
            "is_correct": is_correct,
            "explanation": q["explanation"],
        })
        results["max_score"] += 1
        results["total_score"] += 1 if is_correct else 0

    for i, q in enumerate(longs):
        student_ans = answers.get(f"long_{i}", "")
        grade = grade_long_answer(q["question_text"], q["model_answer"], q.get("key_points", []), student_ans)
        score_percent = grade.get("score_percent", 0)
        results["long"].append({
            "question": q["question_text"],
            "student_answer": student_ans,
            "model_answer": q["model_answer"],
            "score_percent": score_percent,
            "feedback": grade.get("feedback", ""),
        })
        results["max_score"] += 1
        results["total_score"] += score_percent / 100

    return results


def reset_test_state():
    st.session_state.test_mode = False
    st.session_state.test_submitted = False
    st.session_state.test_results = None


# ==========================================
# 3. Page Configuration & Minimalist CSS
# ==========================================
st.set_page_config(page_title="AI Quiz Generator", page_icon="⚡", layout="centered")

st.markdown("""
    <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}

        div.stButton > button:first-child {
            width: 100%;
            border-radius: 6px;
            height: 3.2em;
            font-weight: 600;
            font-size: 1.1rem;
            margin-top: 10px;
        }

        [data-testid="stFileUploadDropzone"] {
            border-radius: 8px;
            padding: 2rem;
        }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 4. Session State Defaults
# ==========================================
for key, default in [
    ("quiz_data", None),
    ("exam_metadata", {}),
    ("test_mode", False),
    ("test_submitted", False),
    ("test_results", None),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ==========================================
# 5. Sidebar (Exam Details + Configuration)
# ==========================================
with st.sidebar:
    st.header("Exam Details")
    st.caption("Shown on the exported paper and the student test screen.")
    institution_name = st.text_input("University / College / School Name")
    department = st.text_input("Department (optional)")
    subject = st.text_input("Subject")
    class_name = st.text_input("Class / Semester")
    teacher_name = st.text_input("Teacher / Examiner Name")
    exam_title = st.text_input("Exam Title", value="AI Generated Quiz")
    duration_minutes = st.number_input("Duration (minutes)", min_value=0, value=60, step=5)
    total_marks = st.number_input("Total Marks", min_value=0, value=100, step=5)

    st.divider()
    st.header("Configuration")
    st.markdown("Choose which question types and how many of each.")

    st.subheader("Multiple Choice")
    num_mcq = st.number_input("Number of MCQs", min_value=0, max_value=100, value=10, step=5)

    st.subheader("Fill in the Blank")
    num_fill_blank = st.number_input("Number of Fill-in-the-Blank", min_value=0, max_value=100, value=0, step=5)

    st.subheader("Long Questions")
    num_long = st.number_input("Number of Long Questions", min_value=0, max_value=50, value=0, step=1)

    difficulty = st.selectbox("Difficulty Level", ["Easy", "Medium", "Hard"])

    st.divider()
    st.caption("Engine: Groq Llama-3.1 8B")
    st.caption("Architecture: Micro-Batched RAG")

exam_metadata = {
    "institution_name": institution_name,
    "department": department,
    "subject": subject,
    "class_name": class_name,
    "teacher_name": teacher_name,
    "exam_title": exam_title,
    "duration_minutes": duration_minutes,
    "total_marks": total_marks,
}
st.session_state.exam_metadata = exam_metadata

# ==========================================
# 6. Main App UI (Header & Uploader)
# ==========================================
st.title("⚡ AI Auto Quiz Generator")
st.markdown("Transform your study materials, books, or notes into comprehensive quizzes — MCQs, fill-in-the-blanks, and long questions.")
st.write("")

uploaded_file = st.file_uploader("Upload PDF Document", type=["pdf"])

# ==========================================
# 7. Generate Logic
# ==========================================
if uploaded_file is not None and not st.session_state.test_mode:
    total_requested = num_mcq + num_fill_blank + num_long

    if st.button("Generate Quiz"):
        if total_requested <= 0:
            st.warning("Please select at least 1 question of some type in the sidebar.")
        else:
            with st.spinner("Analyzing document and generating quiz... This may take a few moments."):
                raw_text = extract_text_from_pdf(uploaded_file)

                if raw_text.startswith("Error"):
                    st.error(raw_text)
                else:
                    try:
                        question_counts = {
                            "mcq": num_mcq,
                            "fill_blank": num_fill_blank,
                            "long_answer": num_long,
                        }
                        st.session_state.quiz_data = generate_quiz_from_large_text(
                            raw_text, question_counts, difficulty
                        )
                        reset_test_state()
                        st.success("Quiz generated successfully!")
                    except Exception as e:
                        st.error(f"An error occurred during generation: {str(e)}")

# ==========================================
# 8. Results & Export Section (Teacher / Preview View)
# ==========================================
if st.session_state.quiz_data and not st.session_state.test_mode:
    quiz_data = st.session_state.quiz_data
    mcqs = quiz_data.get("mcq_questions", [])
    fill_blanks = quiz_data.get("fill_blank_questions", [])
    longs = quiz_data.get("long_questions", [])

    st.divider()

    col1, col2, col3 = st.columns([0.5, 0.25, 0.25])
    with col1:
        st.subheader("Quiz Preview")
    with col2:
        docx_file = create_word_docx(quiz_data, exam_metadata)
        st.download_button(
            label="📄 Export to Word",
            data=docx_file,
            file_name="AI_Quiz_Export.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col3:
        if st.button("🎓 Start Test Mode"):
            st.session_state.test_mode = True
            st.session_state.test_submitted = False
            st.session_state.test_results = None
            st.rerun()

    st.write("")

    if mcqs:
        st.markdown("### Section A: Multiple Choice Questions")
        for i, q in enumerate(mcqs):
            with st.container():
                st.markdown(f"**{i+1}. {q['question_text']}**")
                for j, opt in enumerate(q["options"]):
                    letter = chr(65 + j)
                    st.write(f"&nbsp;&nbsp;&nbsp;&nbsp;**{letter})** {opt}")
                with st.expander("Show Answer"):
                    st.markdown(f"**Correct Answer:** {q['correct_answer']}")
                    st.markdown(f"**Explanation:** *{q['explanation']}*")
                st.write("")

    if fill_blanks:
        st.markdown("### Section B: Fill in the Blanks")
        for i, q in enumerate(fill_blanks):
            with st.container():
                st.markdown(f"**{i+1}. {q['question_text']}**")
                with st.expander("Show Answer"):
                    st.markdown(f"**Correct Answer:** {q['correct_answer']}")
                    st.markdown(f"**Explanation:** *{q['explanation']}*")
                st.write("")

    if longs:
        st.markdown("### Section C: Long Questions")
        for i, q in enumerate(longs):
            with st.container():
                st.markdown(f"**{i+1}. {q['question_text']}**")
                with st.expander("Show Model Answer"):
                    st.markdown(f"**Model Answer:** {q['model_answer']}")
                    if q.get("key_points"):
                        st.markdown("**Key Points:**")
                        for kp in q["key_points"]:
                            st.markdown(f"- {kp}")
                st.write("")

# ==========================================
# 9. Interactive Test-Taking Mode (Student View)
# ==========================================
if st.session_state.test_mode and st.session_state.quiz_data:
    quiz_data = st.session_state.quiz_data
    mcqs = quiz_data.get("mcq_questions", [])
    fill_blanks = quiz_data.get("fill_blank_questions", [])
    longs = quiz_data.get("long_questions", [])

    st.divider()

    if exam_metadata.get("institution_name"):
        st.markdown(f"## {exam_metadata['institution_name']}")
    st.markdown(f"### {exam_metadata.get('exam_title') or 'Exam'}")

    info_bits = []
    if exam_metadata.get("subject"):
        info_bits.append(f"**Subject:** {exam_metadata['subject']}")
    if exam_metadata.get("class_name"):
        info_bits.append(f"**Class:** {exam_metadata['class_name']}")
    if exam_metadata.get("duration_minutes"):
        info_bits.append(f"**Duration:** {exam_metadata['duration_minutes']} mins")
    if exam_metadata.get("total_marks"):
        info_bits.append(f"**Total Marks:** {exam_metadata['total_marks']}")
    if info_bits:
        st.markdown(" &nbsp;|&nbsp; ".join(info_bits))

    st.write("")

    # ---------- Results view (after submission) ----------
    if st.session_state.test_submitted and st.session_state.test_results:
        results = st.session_state.test_results
        pct = (results["total_score"] / results["max_score"] * 100) if results["max_score"] else 0

        st.header("📊 Test Results")
        st.metric("Final Score", f"{results['total_score']:.1f} / {results['max_score']}", f"{pct:.1f}%")
        st.write("")

        if results["mcq"]:
            st.markdown("#### Section A: MCQs")
            for i, item in enumerate(results["mcq"]):
                icon = "✅" if item["is_correct"] else "❌"
                st.markdown(f"{icon} **{i+1}. {item['question']}**")
                st.write(f"Your answer: {item['selected'] or '*Not answered*'}")
                if not item["is_correct"]:
                    st.write(f"Correct answer: {item['correct_answer']}")
                st.caption(item["explanation"])
                st.write("")

        if results["fill_blank"]:
            st.markdown("#### Section B: Fill in the Blanks")
            for i, item in enumerate(results["fill_blank"]):
                icon = "✅" if item["is_correct"] else "❌"
                st.markdown(f"{icon} **{i+1}. {item['question']}**")
                st.write(f"Your answer: {item['student_answer'] or '*Not answered*'}")
                if not item["is_correct"]:
                    st.write(f"Correct answer: {item['correct_answer']}")
                st.caption(item["explanation"])
                st.write("")

        if results["long"]:
            st.markdown("#### Section C: Long Questions")
            for i, item in enumerate(results["long"]):
                st.markdown(f"**{i+1}. {item['question']}**")
                st.write(f"Your answer: {item['student_answer'] or '*Not answered*'}")
                st.write(f"Score: {item['score_percent']}%")
                st.caption(item["feedback"])
                with st.expander("Model Answer"):
                    st.write(item["model_answer"])
                st.write("")

        if st.button("🔄 Back to Quiz Setup"):
            reset_test_state()
            st.rerun()

    # ---------- Test-taking form ----------
    else:
        with st.form("test_form"):
            student_name = st.text_input("Your Name")
            answers = {}
            q_num = 0

            if mcqs:
                st.markdown("#### Section A: Multiple Choice Questions")
                for i, q in enumerate(mcqs):
                    q_num += 1
                    answers[f"mcq_{i}"] = st.radio(
                        f"Q{q_num}. {q['question_text']}",
                        [PLACEHOLDER] + q["options"],
                        key=f"mcq_{i}"
                    )

            if fill_blanks:
                st.markdown("#### Section B: Fill in the Blanks")
                for i, q in enumerate(fill_blanks):
                    q_num += 1
                    answers[f"fb_{i}"] = st.text_input(f"Q{q_num}. {q['question_text']}", key=f"fb_{i}")

            if longs:
                st.markdown("#### Section C: Long Questions")
                for i, q in enumerate(longs):
                    q_num += 1
                    answers[f"long_{i}"] = st.text_area(f"Q{q_num}. {q['question_text']}", key=f"long_{i}")

            submitted = st.form_submit_button("✅ Submit Test")

        if submitted:
            with st.spinner("Grading your answers..."):
                st.session_state.test_results = grade_test(quiz_data, answers)
                st.session_state.test_submitted = True
            st.rerun()

        if st.button("⬅ Exit Test Mode"):
            reset_test_state()
            st.rerun()